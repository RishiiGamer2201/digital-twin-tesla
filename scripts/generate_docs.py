"""Generate project documentation as a .docx file for Google Docs upload."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ---- Styles ----
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.3

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0, 51, 102)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Digital Twin of Nikola Tesla')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Project Documentation')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('AIMS DTU Summer Project 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('June 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_paragraph()
repo = doc.add_paragraph()
repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = repo.add_run('GitHub: github.com/RishiiGamer2201/digital-twin-tesla')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0, 102, 204)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. Problem Statement',
    '3. Objectives',
    '4. System Architecture',
    '5. Technology Stack',
    '6. Features',
    '   6.1 Core AI Features',
    '   6.2 Interactive Features',
    '   6.3 UI/UX Features',
    '7. Data Collection Pipeline',
    '8. RAG Pipeline',
    '9. Memory System',
    '10. Persona Engineering',
    '11. API Endpoints',
    '12. Design Decisions',
    '13. Project Structure',
    '14. Setup and Installation',
    '15. Testing and Sample Conversations',
    '16. Future Scope',
    '17. Conclusion',
    '18. References',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'This project presents a Digital Twin of Nikola Tesla, an AI-powered conversational agent '
    'that simulates the personality, knowledge, and speaking style of the legendary inventor '
    'and electrical engineer (1856-1943). The system uses Retrieval-Augmented Generation (RAG) '
    'to ground Tesla\'s responses in his actual writings, patents, and historical records, '
    'combined with a detailed persona engine that captures his distinctive Victorian-era speech '
    'patterns, Serbian heritage, and visionary philosophy.'
)
doc.add_paragraph(
    'The Digital Twin is not a simple chatbot. It combines multiple AI techniques: '
    'vector similarity search over Tesla\'s actual texts, a two-tier memory system '
    '(short-term conversation buffer and long-term SQLite storage), streaming response '
    'generation via Google Gemini 2.5 Flash, and a rich web-based frontend with interactive '
    'features including a knowledge graph, timeline, quiz mode, and a Tesla vs Edison debate mode.'
)
doc.add_paragraph(
    'The entire system runs at zero cost using free-tier APIs and open-source tools, '
    'making it accessible for educational and research purposes.'
)

# ============================================================
# 2. PROBLEM STATEMENT
# ============================================================
doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph(
    'Historical figures like Nikola Tesla have left behind a wealth of writings, interviews, '
    'and technical documents, but engaging with this material in a conversational, interactive '
    'manner is not possible through traditional means. Students and enthusiasts must read through '
    'dense historical texts to understand Tesla\'s perspectives, with no way to ask follow-up '
    'questions or explore topics dynamically.'
)
doc.add_paragraph(
    'The challenge is to create an AI system that can faithfully represent Tesla\'s personality, '
    'knowledge, and speaking style while being grounded in verifiable historical sources, '
    'maintaining conversation context across sessions, and providing an engaging, modern user experience.'
)

# ============================================================
# 3. OBJECTIVES
# ============================================================
doc.add_heading('3. Objectives', level=1)
objectives = [
    'Build a RAG pipeline that retrieves relevant passages from Tesla\'s actual writings before generating responses.',
    'Create a detailed persona engine that captures Tesla\'s distinctive voice, epistemic style, and timeline awareness.',
    'Implement a two-tier memory system allowing the agent to remember facts across conversation sessions.',
    'Develop a premium web frontend with streaming responses, voice I/O, and interactive visualizations.',
    'Add educational features: quiz mode, knowledge graph, interactive timeline, and debate mode.',
    'Achieve all of the above at zero cost using free-tier services and open-source tools.',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

# ============================================================
# 4. SYSTEM ARCHITECTURE
# ============================================================
doc.add_heading('4. System Architecture', level=1)
doc.add_paragraph(
    'The system follows a layered architecture with clear separation of concerns:'
)

doc.add_heading('4.1 Architecture Overview', level=2)
arch_text = """User (Browser)
    |
    v
[HTML/CSS/JS Frontend] --- Voice I/O (Web Speech API)
    |
    | REST API (JSON) + SSE (streaming)
    v
[Flask Server] -----> [RAG Retriever] -----> [ChromaDB + sentence-transformers]
    |                                                  |
    v                                                  v
[Digital Twin Agent] <--- [Relevant Chunks] <--- [Vector Store]
    |
    v
[Memory Manager] --> [Short-term Buffer (20 msgs)]
    |                 [Long-term SQLite (sessions + facts)]
    v
[Gemini 2.5 Flash] <--- [Tesla Persona Config]
                   <--- [Edison Persona Config] (debate mode)"""
p = doc.add_paragraph(arch_text)
p.style.font.name = 'Consolas'
p.style.font.size = Pt(9)

doc.add_heading('4.2 Data Flow (Streaming Chat)', level=2)
flow_steps = [
    'User types a message in the browser.',
    'Frontend sends POST request to /api/chat/stream.',
    'Server invokes twin.chat_stream(message).',
    'RAG Retriever searches ChromaDB for top-5 relevant chunks.',
    'Memory Manager retrieves short-term history and long-term memories.',
    'Prompt Builder assembles system prompt with persona + RAG context + memory.',
    'Gemini generates response with stream=True, yielding tokens.',
    'Server streams tokens as SSE events to the browser.',
    'After completion: emotion detection, confidence scoring, and suggestion generation.',
    'Frontend renders tokens with typing animation, then shows metadata (badges, sources, chips).',
]
for i, step in enumerate(flow_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

# ============================================================
# 5. TECHNOLOGY STACK
# ============================================================
doc.add_heading('5. Technology Stack', level=1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Component'
hdr[1].text = 'Technology'
hdr[2].text = 'Cost'
for h in hdr:
    for p in h.paragraphs:
        p.runs[0].bold = True

stack = [
    ('LLM', 'Google Gemini 2.5 Flash', 'Free'),
    ('Embeddings', 'sentence-transformers (all-MiniLM-L6-v2, 384 dims)', 'Free'),
    ('Vector Store', 'ChromaDB (local, persistent)', 'Free'),
    ('Long-term Memory', 'SQLite', 'Free'),
    ('Web Server', 'Flask (Python)', 'Free'),
    ('Frontend', 'HTML / CSS / JavaScript', 'Free'),
    ('Voice I/O', 'Web Speech API (browser-native)', 'Free'),
    ('Data Sources', 'Public domain texts, YouTube transcripts', 'Free'),
    ('Total Cost', '', '$0'),
]
for comp, tech, cost in stack:
    row = table.add_row().cells
    row[0].text = comp
    row[1].text = tech
    row[2].text = cost

# ============================================================
# 6. FEATURES
# ============================================================
doc.add_heading('6. Features', level=1)

doc.add_heading('6.1 Core AI Features', level=2)

doc.add_heading('6.1.1 RAG Pipeline (Retrieval-Augmented Generation)', level=3)
doc.add_paragraph(
    'Every response Tesla gives is grounded in real source material. The system retrieves '
    'relevant passages from Tesla\'s autobiography ("My Inventions"), Wikipedia articles, '
    'Wikiquote collections, and YouTube documentary transcripts before generating a response. '
    'This means Tesla\'s answers are backed by actual historical text rather than hallucinated.'
)

doc.add_heading('6.1.2 Deep Persona Engine', level=3)
doc.add_paragraph(
    'Tesla\'s personality is encoded in a 3000+ token system prompt covering his speech '
    'patterns (formal Victorian English with Serbian touches), epistemic style (favoring '
    'mental visualization over brute-force experimentation), teaching approach (vivid analogies '
    'to nature), and timeline awareness (knowledge cutoff at January 1943). The persona includes '
    'his opinions on Edison, his friendship with Mark Twain, his love for pigeons, and his '
    'passion for free wireless energy.'
)

doc.add_heading('6.1.3 Two-Tier Memory System', level=3)
doc.add_paragraph(
    'Short-term memory: A rolling buffer of the last 20 messages provides immediate context '
    'within a session. Long-term memory: When a session ends, Gemini generates a structured '
    'summary (topics, key facts, entities) stored in SQLite. In future sessions, relevant past '
    'memories are retrieved and injected into context. Tell Tesla your name once, and he '
    'remembers it next time.'
)

doc.add_heading('6.1.4 Source Citations', level=3)
doc.add_paragraph(
    'Every response shows which Tesla writings were referenced, displayed as clickable chips. '
    'Users can expand a preview of the actual source text to verify historical grounding.'
)

doc.add_heading('6.1.5 Fact Confidence Badges', level=3)
doc.add_paragraph(
    'Each response is tagged with a confidence level based on RAG retrieval scores: '
    '"From my writings" (green, high similarity), "From my recollection" (amber, partial match), '
    'or "Beyond my time" (red, speculative/post-1943 topics).'
)

doc.add_heading('6.2 Interactive Features', level=2)

doc.add_heading('6.2.1 Streaming Responses with Typing Animation', level=3)
doc.add_paragraph(
    'Responses stream token-by-token from the Gemini API via Server-Sent Events (SSE). '
    'Text appears character-by-character in real time, creating a natural typing effect. '
    'The backend uses Flask\'s Response generator with text/event-stream MIME type. The '
    'frontend reads the stream using the Fetch API\'s ReadableStream reader.'
)

doc.add_heading('6.2.2 Topic Suggestion Chips', level=3)
doc.add_paragraph(
    'After each response, 2-3 clickable follow-up question chips appear, generated by Gemini '
    'based on the conversation context. Clicking a chip sends that question automatically.'
)

doc.add_heading('6.2.3 Emotion Indicator', level=3)
doc.add_paragraph(
    'A colored dot on Tesla\'s avatar changes based on his emotional tone: '
    'gold (excited, discussing inventions), red (bitter, discussing Edison/business), '
    'purple (nostalgic, recalling childhood/Serbia), blue (philosophical, discussing the universe), '
    'cyan (neutral, default). Detected by a separate Gemini classification call.'
)

doc.add_heading('6.2.4 Interactive Timeline (1856-1943)', level=3)
doc.add_paragraph(
    'A visual vertical timeline showing 16 key events in Tesla\'s life. Features a glowing '
    'connecting line, hover effects on event cards, and click-to-ask functionality that sends '
    'the event as a question to the chat.'
)

doc.add_heading('6.2.5 Knowledge Graph Visualization', level=3)
doc.add_paragraph(
    'An interactive canvas-based graph with approximately 18 nodes showing connections between Tesla\'s '
    'inventions (cyan), people (gold), places (purple), and concepts (green). Uses a simple '
    'force-directed layout algorithm. Click any node to ask Tesla about that topic.'
)

doc.add_heading('6.2.6 Quiz Mode', level=3)
doc.add_paragraph(
    'Interactive multiple-choice quiz with questions generated by Gemini on demand. Features '
    '4 options per question, visual feedback (green/red highlighting), Tesla-style grading, '
    'running score tracker, and unlimited questions via "Next Question" button.'
)

doc.add_heading('6.2.7 Tesla vs Edison Debate Mode', level=3)
doc.add_paragraph(
    'A split-panel debate interface where Tesla and Edison argue their positions on any topic. '
    'Both responses are generated by Gemini using separate, historically accurate persona prompts. '
    'Edison\'s persona is practical, business-minded, and favors brute-force experimentation.'
)

doc.add_heading('6.2.8 PDF Upload', level=3)
doc.add_paragraph(
    'Users can upload Tesla-related PDFs to expand the knowledge base at runtime. The system '
    'extracts text using pypdf, chunks it into 512-word segments, generates embeddings, and '
    'adds them to ChromaDB. Future responses will reference the uploaded material.'
)

doc.add_heading('6.2.9 Past Session Revisit', level=3)
doc.add_paragraph(
    'Users can click any past session in the sidebar to view its full details: summary, '
    'topics discussed, key entities, and remembered facts. A "Continue This Conversation" '
    'button resumes the chat with the session\'s context.'
)

doc.add_heading('6.3 UI/UX Features', level=2)

doc.add_heading('6.3.1 Dark/Light Theme Toggle', level=3)
doc.add_paragraph(
    'Toggle between dark mode (electric blue on deep navy) and light mode (blue on white). '
    'Theme preference persists across sessions via localStorage. Smooth CSS transitions.'
)

doc.add_heading('6.3.2 Mobile Responsive Design', level=3)
doc.add_paragraph(
    'On screens smaller than 768px, the sidebar collapses and a hamburger menu appears. '
    'All features remain fully accessible on mobile devices.'
)

doc.add_heading('6.3.3 Voice Input and Output', level=3)
doc.add_paragraph(
    'Voice output uses the browser\'s Speech Synthesis API. Voice input uses the Web Speech API '
    'for speech recognition. No server-side dependencies; everything runs in the browser.'
)

doc.add_heading('6.3.4 Conversation Export', level=3)
doc.add_paragraph(
    'Download the entire chat as a formatted Markdown file with timestamped export, '
    'role labels, and clean formatting.'
)

# ============================================================
# 7. DATA COLLECTION
# ============================================================
doc.add_heading('7. Data Collection Pipeline', level=1)
doc.add_paragraph('The system collects Tesla\'s writings from multiple public domain sources:')

sources_table = doc.add_table(rows=1, cols=3)
sources_table.style = 'Light Grid Accent 1'
hdr = sources_table.rows[0].cells
hdr[0].text = 'Source'
hdr[1].text = 'Content'
hdr[2].text = 'Method'
for h in hdr:
    for p in h.paragraphs:
        p.runs[0].bold = True

sources_data = [
    ('Project Gutenberg', '"My Inventions" autobiography (1919)', 'Web scraping (BeautifulSoup)'),
    ('Wikipedia', 'Biographical and technical articles', 'Web scraping'),
    ('Wikiquote', '200+ verified Tesla quotes', 'Web scraping'),
    ('YouTube', '8 documentary transcripts', 'youtube-transcript-api'),
    ('User uploads', 'Additional PDFs (at runtime)', 'pypdf text extraction'),
]
for src, content, method in sources_data:
    row = sources_table.add_row().cells
    row[0].text = src
    row[1].text = content
    row[2].text = method

doc.add_paragraph()
doc.add_paragraph(
    'All raw data is processed through a preprocessing pipeline that cleans text, '
    'removes noise, and splits into 512-word chunks with 64-word overlap. Each chunk '
    'carries metadata: source name, topic, year, and scientist ("Nikola Tesla"). '
    'The final corpus contains approximately 188 document chunks.'
)

# ============================================================
# 8. RAG PIPELINE
# ============================================================
doc.add_heading('8. RAG Pipeline', level=1)
doc.add_paragraph(
    'The Retrieval-Augmented Generation pipeline ensures Tesla\'s responses are grounded in facts:'
)
rag_steps = [
    'Embedder: Uses all-MiniLM-L6-v2 (384 dimensions) from sentence-transformers for fast, '
    'local embedding generation. No API calls needed.',
    'Vector Store: ChromaDB with cosine similarity, persistent local storage in .chroma/ directory. '
    'Supports runtime expansion via PDF uploads.',
    'Retriever: Returns top-5 relevant chunks with formatted source labels. Each chunk includes '
    'the source work, chapter, and year for citation.',
    'Confidence Scoring: Average cosine similarity across retrieved chunks determines confidence '
    'level (grounded > 0.5, inferred 0.3-0.5, speculative < 0.3).',
]
for step in rag_steps:
    doc.add_paragraph(step, style='List Bullet')

# ============================================================
# 9. MEMORY SYSTEM
# ============================================================
doc.add_heading('9. Memory System', level=1)

doc.add_heading('9.1 Short-Term Memory', level=2)
doc.add_paragraph(
    'A rolling buffer of the last 20 messages within a session. Provides immediate '
    'conversational context to the LLM. Reset when a session ends. Formatted as '
    '"Visitor/Tesla" pairs for prompt injection.'
)

doc.add_heading('9.2 Long-Term Memory', level=2)
doc.add_paragraph(
    'SQLite database with two tables: sessions (id, timestamp, summary, topics, entities, '
    'message_count, user_name) and facts (id, session_id, content, category, importance, '
    'created_at). When a session ends, Gemini generates a structured summary and extracts '
    'key facts. In future sessions, relevant past memories are retrieved by keyword matching '
    'and injected into Tesla\'s context.'
)

# ============================================================
# 10. PERSONA ENGINEERING
# ============================================================
doc.add_heading('10. Persona Engineering', level=1)

doc.add_heading('10.1 Tesla Persona', level=2)
doc.add_paragraph(
    'The Tesla persona is defined in a 3000+ token system prompt covering:'
)
persona_items = [
    'Voice and Speech Patterns: Formal Victorian-era English with occasional Serbian expressions. '
    'Precise, measured language. Vivid visual descriptions leveraging his eidetic memory.',
    'Epistemic Style: Deep belief in mental visualization and elegant mathematical solutions. '
    'Distrust of Edison\'s brute-force approach. The universe as energy, frequency, and vibration.',
    'Teaching Style: Explanations through vivid mental pictures and analogies to nature. '
    'Starting from fundamental principles of electromagnetism.',
    'Timeline Awareness: Knowledge extends to January 1943. For post-1943 topics, Tesla says '
    '"That would be after my time, but I can tell you what I envisioned..."',
    'Core Areas: AC power, Tesla coil, rotating magnetic field, radio, wireless energy, '
    'X-rays, remote control, hydroelectric power.',
    'Personal Details: References to Smiljan, Graz, Colorado Springs, Wardenclyffe. '
    'Friendship with Mark Twain. Love for pigeons. Frustration with commerce corrupting invention.',
]
for item in persona_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10.2 Edison Persona (Debate Mode)', level=2)
doc.add_paragraph(
    'A historically accurate Edison persona used only in debate mode. Edison is practical, '
    'business-minded, direct, and speaks in plain American English. He defends DC power, '
    'values perspiration over inspiration, and is skeptical of Tesla\'s grand visions.'
)

# ============================================================
# 11. API ENDPOINTS
# ============================================================
doc.add_heading('11. API Endpoints', level=1)

api_table = doc.add_table(rows=1, cols=3)
api_table.style = 'Light Grid Accent 1'
hdr = api_table.rows[0].cells
hdr[0].text = 'Method'
hdr[1].text = 'Endpoint'
hdr[2].text = 'Description'
for h in hdr:
    for p in h.paragraphs:
        p.runs[0].bold = True

apis = [
    ('GET', '/', 'Serve the frontend'),
    ('POST', '/api/chat', 'Non-streaming chat (fallback)'),
    ('POST', '/api/chat/stream', 'Streaming chat via SSE'),
    ('GET', '/api/memory', 'Memory dashboard statistics'),
    ('POST', '/api/session/end', 'End session, save to long-term memory'),
    ('POST', '/api/session/clear', 'Clear chat without saving'),
    ('POST', '/api/session/delete', 'Delete a specific past session'),
    ('POST', '/api/session/details', 'Get full details of a past session'),
    ('POST', '/api/quiz/generate', 'Generate a quiz question'),
    ('POST', '/api/quiz/check', 'Grade a quiz answer'),
    ('POST', '/api/debate', 'Tesla vs Edison debate'),
    ('POST', '/api/export', 'Export chat as Markdown'),
    ('GET', '/api/timeline', 'Tesla life timeline events'),
    ('POST', '/api/upload-pdf', 'Upload PDF to knowledge base'),
]
for method, endpoint, desc in apis:
    row = api_table.add_row().cells
    row[0].text = method
    row[1].text = endpoint
    row[2].text = desc

# ============================================================
# 12. DESIGN DECISIONS
# ============================================================
doc.add_heading('12. Key Design Decisions', level=1)

decisions = [
    ('Flask over Streamlit',
     'Flask provides full control over streaming (SSE), modals, canvas graphics, and '
     'browser-native voice. Streamlit\'s rerun model was problematic for real-time features.'),
    ('SSE over WebSockets',
     'Server-Sent Events are simpler, require no additional dependencies, and are ideal '
     'for unidirectional token streaming. WebSockets would add unnecessary complexity.'),
    ('sentence-transformers over API Embeddings',
     'Local embeddings provide zero cost, complete privacy, and fast performance (~50ms per batch). '
     'Sufficient quality for retrieval with 384 dimensions.'),
    ('ChromaDB over Pinecone/Weaviate',
     'Zero cost, zero config, Python-native, persistent, and sufficient scale for our corpus. '
     'Supports runtime expansion via PDF uploads.'),
    ('System Prompt over Fine-tuning',
     'No training data needed. Immediate iteration by editing text. Sufficient quality with '
     'modern LLMs. Transparent persona definition readable in tesla_config.py.'),
    ('Canvas over D3.js for Knowledge Graph',
     'Zero dependencies, simpler code, sufficient for approximately 20 nodes. No external library to download.'),
    ('Web Speech API over Server-side Voice',
     'Browser-native, instant latency, zero dependencies. Eliminates audio format conversion '
     'issues that plagued the original server-side approach.'),
]
for title, rationale in decisions:
    doc.add_heading(title, level=3)
    doc.add_paragraph(rationale)

# ============================================================
# 13. PROJECT STRUCTURE
# ============================================================
doc.add_heading('13. Project Structure', level=1)
structure = """digital-twin-tesla/
|-- .env.example                # Environment variable template
|-- .gitignore                  # Git ignore rules
|-- README.md                   # Project documentation
|-- requirements.txt            # Python dependencies
|-- setup.py                    # Package setup
|-- server.py                   # Flask backend (14 API endpoints)
|-- static/
|   |-- index.html              # Complete frontend (HTML/CSS/JS)
|-- src/
|   |-- __init__.py
|   |-- agent/
|   |   |-- __init__.py
|   |   |-- twin.py             # Core Digital Twin agent
|   |-- data_collection/
|   |   |-- __init__.py
|   |   |-- scraper.py          # Web scraper for Tesla texts
|   |   |-- youtube_loader.py   # YouTube transcript downloader
|   |   |-- pdf_loader.py       # PDF text extractor
|   |   |-- preprocessor.py     # Text chunking and metadata
|   |-- llm/
|   |   |-- __init__.py
|   |   |-- gemini_client.py    # Gemini API wrapper
|   |-- memory/
|   |   |-- __init__.py
|   |   |-- short_term.py       # Conversation buffer
|   |   |-- long_term.py        # SQLite session/fact storage
|   |   |-- manager.py          # Unified memory orchestrator
|   |-- persona/
|   |   |-- __init__.py
|   |   |-- tesla_config.py     # Tesla persona (system prompt)
|   |   |-- edison_config.py    # Edison persona (debate mode)
|   |   |-- prompt_builder.py   # Dynamic prompt assembly
|   |-- rag/
|       |-- __init__.py
|       |-- embedder.py         # Local sentence-transformer embeddings
|       |-- vectorstore.py      # ChromaDB interface
|       |-- retriever.py        # High-level retrieval logic
|-- scripts/
|   |-- collect_data.py         # Data collection runner
|   |-- build_vectorstore.py    # Vector store builder
|-- docs/
|   |-- architecture.md         # System architecture diagrams
|   |-- design_decisions.md     # Technical design rationale
|-- tests/
|   |-- sample_conversations.md # Example conversations for testing"""
p = doc.add_paragraph(structure)
p.style.font.size = Pt(9)

# ============================================================
# 14. SETUP
# ============================================================
doc.add_heading('14. Setup and Installation', level=1)
setup_steps = [
    'Clone: git clone https://github.com/RishiiGamer2201/digital-twin-tesla.git',
    'Virtual environment: python -m venv venv && venv\\Scripts\\activate',
    'Install dependencies: pip install -r requirements.txt',
    'Environment: Copy .env.example to .env and add Gemini API key from aistudio.google.com',
    'Data collection (one-time): python scripts/collect_data.py',
    'Build vector store (one-time): python scripts/build_vectorstore.py',
    'Run: python server.py (opens at http://localhost:5000)',
]
for i, step in enumerate(setup_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

# ============================================================
# 15. TESTING
# ============================================================
doc.add_heading('15. Testing and Sample Conversations', level=1)
doc.add_paragraph('The following questions test different system capabilities:')

test_table = doc.add_table(rows=1, cols=3)
test_table.style = 'Light Grid Accent 1'
hdr = test_table.rows[0].cells
hdr[0].text = 'Category'
hdr[1].text = 'Question'
hdr[2].text = 'Tests'
for h in hdr:
    for p in h.paragraphs:
        p.runs[0].bold = True

tests = [
    ('Persona', 'Tell me about your autobiography My Inventions', 'RAG retrieval + voice'),
    ('Persona', 'What do you think about Thomas Edison?', 'Persona opinions'),
    ('Technical', 'How does a Tesla coil work?', 'Technical depth + RAG'),
    ('Technical', 'Why is AC better than DC for long distances?', 'Core expertise'),
    ('Timeline', 'What do you think about the internet?', 'Post-1943 handling'),
    ('Memory', 'My name is Sagar, I study at DTU', 'Fact extraction + recall'),
    ('Debate', 'AC vs DC power (in debate mode)', 'Dual persona generation'),
    ('Quiz', 'Any quiz question', 'Question generation + grading'),
]
for cat, q, t in tests:
    row = test_table.add_row().cells
    row[0].text = cat
    row[1].text = q
    row[2].text = t

# ============================================================
# 16. FUTURE SCOPE
# ============================================================
doc.add_heading('16. Future Scope', level=1)
future = [
    'Multi-language support (Tesla spoke Serbian, German, French, and English).',
    'Adding more historical figures for cross-era conversations.',
    'Fine-tuning a smaller model on Tesla\'s actual writing style for offline use.',
    'Deploying as a public web application using Docker and cloud hosting.',
    'Adding image generation to visualize Tesla\'s inventions.',
    'Integrating with educational platforms for classroom use.',
]
for f in future:
    doc.add_paragraph(f, style='List Bullet')

# ============================================================
# 18. CONCLUSION
# ============================================================
doc.add_heading('17. Conclusion', level=1)
doc.add_paragraph(
    'The Digital Twin of Nikola Tesla demonstrates how modern AI techniques, specifically '
    'Retrieval-Augmented Generation, persona engineering, and long-term memory, can be combined '
    'to create an engaging, educational conversational agent that faithfully represents a '
    'historical figure. The system is grounded in Tesla\'s actual writings, maintains context '
    'across sessions, and provides a rich interactive experience through streaming responses, '
    'quizzes, debates, and visualizations.'
)
doc.add_paragraph(
    'By using exclusively free-tier services and open-source tools, the project proves that '
    'sophisticated AI applications can be built at zero cost, making them accessible for '
    'educational institutions and individual researchers.'
)

# ============================================================
# 19. REFERENCES
# ============================================================
doc.add_heading('18. References', level=1)
refs = [
    'Tesla, Nikola. "My Inventions: The Autobiography of Nikola Tesla." 1919. (Project Gutenberg)',
    'Google. "Gemini API Documentation." aistudio.google.com',
    'Hugging Face. "sentence-transformers/all-MiniLM-L6-v2." huggingface.co',
    'ChromaDB. "ChromaDB Documentation." docs.trychroma.com',
    'Wikipedia. "Nikola Tesla." en.wikipedia.org/wiki/Nikola_Tesla',
    'Wikiquote. "Nikola Tesla." en.wikiquote.org/wiki/Nikola_Tesla',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'[{i}] {ref}')

# ---- Save ----
output_path = r'C:\Users\seast\VSCODE\digital-twin-tesla\docs\Digital_Twin_Tesla_Documentation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
